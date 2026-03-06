"""
Build DOCX template for "Surat Permohonan Magang Berdampak" STIH Adhyaksa.
Matches the PDF reference: 099_Surat Permohonan Magang Berdampak_Simon...

PhpWord placeholders (${...}) are used for all dynamic fields so
InternshipLetterService::generateRequestLetter() can fill them in.

Usage:
    python scripts/build_surat_magang_template.py
Output:
    docs/Surat Permohonan Magang.docx
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTPUT   = os.path.join(BASE_DIR, "docs", "Surat Permohonan Magang.docx")
LOGO     = os.path.join(BASE_DIR, "public", "images", "logo_stih.png")
# Fallback logos if main doesn't exist
LOGO_FALLBACKS = [
    os.path.join(BASE_DIR, "public", "images", "logo_stih_white-clear.png"),
    os.path.join(BASE_DIR, "public", "images", "logo_stih_white.png"),
    os.path.join(BASE_DIR, "public", "images", "certified.png"),
]

# ── Helpers ───────────────────────────────────────────────────────────────────

def find_logo():
    if os.path.exists(LOGO):
        return LOGO
    for f in LOGO_FALLBACKS:
        if os.path.exists(f):
            return f
    return None

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    el = OxmlElement('w:tcBorders')
    for edge, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right),
                      ('insideH', None), ('insideV', None)]:
        e = OxmlElement(f'w:{edge}')
        if val:
            e.set(qn('w:val'), val.get('val', 'single'))
            e.set(qn('w:sz'), str(val.get('sz', 4)))
            e.set(qn('w:space'), '0')
            e.set(qn('w:color'), val.get('color', '000000'))
        else:
            e.set(qn('w:val'), 'nil')
        el.append(e)
    tcPr.append(el)

def set_run_font(run, name='Times New Roman', size_pt=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Force font for non-latin too
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)
    existing = rPr.find(qn('w:rFonts'))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)

def para_fmt(p, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=0,
             line_rule=WD_LINE_SPACING.EXACTLY, line_pt=None, first_indent=None, left_indent=None):
    pf = p.paragraph_format
    pf.alignment = align
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line_pt:
        pf.line_spacing_rule = line_rule
        pf.line_spacing = Pt(line_pt)
    if first_indent is not None:
        pf.first_line_indent = Cm(first_indent)
    if left_indent is not None:
        pf.left_indent = Cm(left_indent)

def add_run(para, text, font='Times New Roman', size=11, bold=False, italic=False):
    r = para.add_run(text)
    set_run_font(r, font, size, bold, italic)
    return r

def add_blank(doc, before=0, after=0):
    p = doc.add_paragraph()
    para_fmt(p, before=before, after=after, line_pt=12)
    return p

def remove_all_borders_from_table(table):
    """Remove all cell borders from a table (used for KOP table)."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    # Kill table-level borders
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'none')
        tblBorders.append(e)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)

def set_table_width(table, width_cm):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(int(width_cm * 567)))  # 1cm = 567 twips
    tblW.set(qn('w:type'), 'dxa')
    existing = tblPr.find(qn('w:tblW'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblW)

def set_col_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_cm * 567)))
    tcW.set(qn('w:type'), 'dxa')
    existing = tcPr.find(qn('w:tcW'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcW)

# ── Build ─────────────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Default paragraph style
    n_style = doc.styles['Normal']
    n_style.font.name = 'Times New Roman'
    n_style.font.size = Pt(11)
    for pf_attr in ['space_before', 'space_after']:
        setattr(n_style.paragraph_format, pf_attr, Pt(0))

    # Page margins (matching reference: top kop surat + margins)
    sec = doc.sections[0]
    sec.page_width  = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin    = Cm(1.5)
    sec.bottom_margin = Cm(1.5)
    sec.left_margin   = Cm(3.0)
    sec.right_margin  = Cm(2.0)

    # Content width = 21 - 3 - 2 = 16 cm
    CONTENT_W = 16.0

    # ── KOP SURAT ─────────────────────────────────────────────────────────────
    # Table: [Logo 3.5cm] | [School info 12.5cm]
    LOGO_W  = 3.5
    INFO_W  = CONTENT_W - LOGO_W  # 12.5cm

    kop = doc.add_table(rows=1, cols=2)
    kop.style = 'Table Grid'
    remove_all_borders_from_table(kop)
    set_table_width(kop, CONTENT_W)

    logo_cell = kop.rows[0].cells[0]
    info_cell = kop.rows[0].cells[1]

    set_col_width(logo_cell, LOGO_W)
    set_col_width(info_cell, INFO_W)

    # Vertical align middle for both cells
    logo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    info_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Only bottom border on entire KOP (thick double line)
    for cell in [logo_cell, info_cell]:
        set_cell_borders(cell,
            top=None, left=None, right=None,
            bottom={'val': 'double', 'sz': 6, 'color': '000000'},
        )

    # Logo cell content
    logo_path = find_logo()
    logo_p = logo_cell.paragraphs[0]
    para_fmt(logo_p, align=WD_ALIGN_PARAGRAPH.CENTER, before=4, after=4)
    if logo_path:
        run = logo_p.add_run()
        run.add_picture(logo_path, height=Cm(2.8))
    else:
        # Placeholder text if no logo
        r = add_run(logo_p, '[LOGO]', size=9)
        r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # School info cell content
    # Line 1: SEKOLAH TINGGI ILMU HUKUM (STIH) ADHYAKSA (large bold, centered)
    p1 = info_cell.paragraphs[0]
    para_fmt(p1, align=WD_ALIGN_PARAGRAPH.CENTER, before=4, after=2, line_pt=14)
    add_run(p1, 'SEKOLAH TINGGI ILMU HUKUM (STIH) ADHYAKSA', size=13, bold=True)

    # Line 2: Jl. address
    p2 = info_cell.add_paragraph()
    para_fmt(p2, align=WD_ALIGN_PARAGRAPH.CENTER, before=1, after=1, line_pt=12)
    add_run(p2, 'Jl. Margasatwa No. 39, Jagakarsa, Kota Jakarta Selatan, DKI Jakarta', size=9)

    # Line 3: Telepon + Email
    p3 = info_cell.add_paragraph()
    para_fmt(p3, align=WD_ALIGN_PARAGRAPH.CENTER, before=1, after=4, line_pt=12)
    add_run(p3, 'Telepon : (021) 220 99999  /  Email : info@stih-adhyaksa.ac.id', size=9)

    # ── BLANK AFTER KOP ───────────────────────────────────────────────────────
    add_blank(doc, before=6, after=0)

    # ── DATE (right-aligned) ──────────────────────────────────────────────────
    p_date = doc.add_paragraph()
    para_fmt(p_date, align=WD_ALIGN_PARAGRAPH.RIGHT, before=0, after=4, line_pt=14)
    add_run(p_date, 'Jakarta,  ')
    add_run(p_date, '${tanggal}')

    # ── NOMOR / LAMPIRAN / PERIHAL block ─────────────────────────────────────
    # Use a table so we can align colon neatly
    meta = doc.add_table(rows=3, cols=3)
    meta.style = 'Table Grid'
    remove_all_borders_from_table(meta)
    set_table_width(meta, CONTENT_W)

    META_LABEL = 3.0
    META_SEP   = 0.5
    META_VAL   = CONTENT_W - META_LABEL - META_SEP

    for row in meta.rows:
        set_col_width(row.cells[0], META_LABEL)
        set_col_width(row.cells[1], META_SEP)
        set_col_width(row.cells[2], META_VAL)
        for cell in row.cells:
            set_cell_borders(cell)  # no borders

    def meta_row(row_idx, label, value):
        r = meta.rows[row_idx]
        p0 = r.cells[0].paragraphs[0]
        para_fmt(p0, before=0, after=0, line_pt=16)
        add_run(p0, label)
        p1 = r.cells[1].paragraphs[0]
        para_fmt(p1, before=0, after=0, line_pt=16)
        add_run(p1, ':')
        p2 = r.cells[2].paragraphs[0]
        para_fmt(p2, before=0, after=0, line_pt=16)
        add_run(p2, value)

    meta_row(0, 'Nomor',    '${nomor_surat}')
    meta_row(1, 'Lampiran', '-')
    meta_row(2, 'Perihal',  'Permohonan Magang Berdampak')

    add_blank(doc, before=4, after=0)

    # ── KEPADA YTH ────────────────────────────────────────────────────────────
    def left_para(text, size=11, bold=False, before=0, after=0, line_pt=16):
        p = doc.add_paragraph()
        para_fmt(p, before=before, after=after, line_pt=line_pt)
        add_run(p, text, size=size, bold=bold)
        return p

    left_para('Kepada Yth.')
    left_para('${nama_pimpinan_instansi}')
    left_para('Di')
    left_para('\u00a0Tempat')   # non-breaking space indent

    add_blank(doc, before=4, after=0)
    left_para('Dengan Hormat,')
    add_blank(doc, before=2, after=0)

    # ── BODY PARAGRAPH 1 ─────────────────────────────────────────────────────
    p_body1 = doc.add_paragraph()
    para_fmt(p_body1, before=0, after=6, line_pt=22, first_indent=1.25)
    add_run(p_body1,
        'Dalam rangka mendukung kegiatan Magang Berdampak dengan ini kami bermaksud '
        'mengajukan permohonan kesediaan Bapak/Ibu untuk bisa menerima mahasiswa Sekolah '
        'Tinggi Ilmu Hukum Adhyaksa untuk melakukan magang di '
    )
    add_run(p_body1, '${instansi}')
    add_run(p_body1,
        ' untuk semester ${tahun_ajaran} yang akan dimulai pada '
    )
    add_run(p_body1, '${periode_mulai}')
    add_run(p_body1, ' sampai dengan ')
    add_run(p_body1, '${periode_selesai}')
    add_run(p_body1, ', adapun mahasiswa yang akan melaksanakan kegiatan tersebut adalah:')

    add_blank(doc, before=2, after=0)

    # ── STUDENT DATA TABLE ────────────────────────────────────────────────────
    mhs = doc.add_table(rows=5, cols=3)
    mhs.style = 'Table Grid'
    remove_all_borders_from_table(mhs)
    set_table_width(mhs, CONTENT_W)

    MHS_INDENT = 4.5   # indent from left margin
    MHS_LABEL  = 4.0   # label column
    MHS_SEP    = 0.5   # colon column
    MHS_VAL    = CONTENT_W - MHS_INDENT - MHS_LABEL - MHS_SEP

    for row in mhs.rows:
        # Use a merged first column for indent effect via padding
        set_col_width(row.cells[0], MHS_LABEL)
        set_col_width(row.cells[1], MHS_SEP)
        set_col_width(row.cells[2], MHS_VAL)
        for cell in row.cells:
            set_cell_borders(cell)
            # Left padding for indent
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            left_mar = OxmlElement('w:left')
            left_mar.set(qn('w:w'), str(int(MHS_INDENT * 567 / len(row.cells) * 0)))
            left_mar.set(qn('w:type'), 'dxa')
            tcMar.append(left_mar)
            tcPr.append(tcMar)

    def mhs_row(row_idx, label, value):
        r = mhs.rows[row_idx]
        for cell in r.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            lm = OxmlElement('w:left')
            # Add left indent to first cell only
            lm.set(qn('w:w'), str(int(MHS_INDENT * 567)) if cell == r.cells[0] else '0')
            lm.set(qn('w:type'), 'dxa')
            tcMar.append(lm)
            existing = tcPr.find(qn('w:tcMar'))
            if existing is not None:
                tcPr.remove(existing)
            tcPr.append(tcMar)

        p0 = r.cells[0].paragraphs[0]
        para_fmt(p0, before=0, after=0, line_pt=18)
        add_run(p0, label)
        p1 = r.cells[1].paragraphs[0]
        para_fmt(p1, before=0, after=0, line_pt=18)
        add_run(p1, ':')
        p2 = r.cells[2].paragraphs[0]
        para_fmt(p2, before=0, after=0, line_pt=18)
        add_run(p2, value)

    mhs_row(0, 'Nama',         '${nama}')
    mhs_row(1, 'NIM',          '${nim}')
    mhs_row(2, 'Semester',     '${semester}')
    mhs_row(3, 'No.Telephone', '${no_hp}')
    mhs_row(4, 'Email',        '${email}')

    add_blank(doc, before=6, after=0)

    # ── BODY PARAGRAPH 2 (Closing) ────────────────────────────────────────────
    p_body2 = doc.add_paragraph()
    para_fmt(p_body2, before=0, after=6, line_pt=22, first_indent=1.25)
    add_run(p_body2,
        'Demikian surat permohonan ini kami sampaikan untuk dapat dipergunakan sebagaimana '
        'mestinya. Atas perhatian dan kerjasamanya, kami ucapkan terima kasih.'
    )

    add_blank(doc, before=6, after=0)

    # ── SIGNATURE BLOCK (right-aligned) ───────────────────────────────────────
    def sig_para(text, bold=False, before=0, after=0, line_pt=16):
        p = doc.add_paragraph()
        para_fmt(p, align=WD_ALIGN_PARAGRAPH.RIGHT, before=before, after=after, line_pt=line_pt)
        add_run(p, text, bold=bold)
        return p

    sig_para('Sekolah Tinggi Ilmu Hukum')
    sig_para('(STIH) Adhyaksa,')
    sig_para('Kepala Bagian Akademik,')

    # Signature space (4 blank lines)
    for _ in range(4):
        p = doc.add_paragraph()
        para_fmt(p, align=WD_ALIGN_PARAGRAPH.RIGHT, before=0, after=0, line_pt=14)

    sig_para('Akhmaad Ikraam, S.H., M.H.', bold=False)

    add_blank(doc, before=8, after=0)

    # ── TEMBUSAN ──────────────────────────────────────────────────────────────
    p_tmb = doc.add_paragraph()
    para_fmt(p_tmb, before=0, after=2, line_pt=16)
    add_run(p_tmb, 'Tembusan :')

    for no, item in enumerate([
        'Ketua STIH Adhyaksa;',
        'Wakil Ketua I STIH Adhyaksa;',
        'Arsip.',
    ], 1):
        p = doc.add_paragraph()
        para_fmt(p, before=0, after=0, line_pt=16, left_indent=0.8, first_indent=-0.8)
        add_run(p, f'{no}. {item}')

    # ── SAVE ─────────────────────────────────────────────────────────────────
    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    doc.save(OUTPUT)
    print(f'✓  Saved to: {OUTPUT}')


if __name__ == '__main__':
    build()
